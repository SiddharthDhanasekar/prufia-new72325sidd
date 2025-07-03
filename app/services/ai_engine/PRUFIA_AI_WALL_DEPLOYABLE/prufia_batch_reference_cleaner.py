
import os
import re
import json
from docx import Document

def exclude_references_and_notes(text):
    """
    Excludes APA, MLA, Chicago-style references and footnote sections from text.
    Recognizes:
    - "References", "Works Cited", "Bibliography"
    - "Notes" or numbered footnote sections (Chicago)
    """
    lines = text.strip().splitlines()

    reference_headers = [
        r'^references\s*$', r'^works cited\s*$', r'^bibliography\s*$',
        r'^citations\s*$', r'^sources\s*$', r'^notes\s*$'
    ]
    header_patterns = [re.compile(pattern, re.IGNORECASE) for pattern in reference_headers]

    start_index = None
    for i, line in enumerate(lines):
        if any(pat.match(line.strip()) for pat in header_patterns):
            start_index = i
            break

    if start_index is not None:
        return '\n'.join(lines[:start_index]).strip()
    return text.strip()

def process_docx_file(input_path):
    """
    Reads a .docx file, extracts text, excludes references/notes, and returns cleaned result.
    """
    doc = Document(input_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    cleaned_text = exclude_references_and_notes(full_text)

    return {
        "filename": os.path.basename(input_path),
        "cleaned_text": cleaned_text,
        "status": "references_removed"
    }

def process_batch_folder(input_folder, output_folder_json=None, output_folder_docx=None):
    """
    Processes all .docx files in input_folder.
    Saves:
    - Cleaned .json per file (if output_folder_json specified)
    - Cleaned .docx per file (if output_folder_docx specified)
    """
    os.makedirs(output_folder_json or "", exist_ok=True)
    os.makedirs(output_folder_docx or "", exist_ok=True)

    for filename in os.listdir(input_folder):
        if filename.lower().endswith(".docx"):
            path = os.path.join(input_folder, filename)
            result = process_docx_file(path)

            # Save JSON
            if output_folder_json:
                json_path = os.path.join(output_folder_json, filename.replace(".docx", ".json"))
                with open(json_path, "w") as f:
                    json.dump(result, f, indent=2)

            # Save cleaned DOCX
            if output_folder_docx:
                doc = Document()
                for line in result["cleaned_text"].splitlines():
                    doc.add_paragraph(line)
                cleaned_docx_path = os.path.join(output_folder_docx, filename)
                doc.save(cleaned_docx_path)

    return "Batch processing complete."
